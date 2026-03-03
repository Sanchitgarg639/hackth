from docx import Document
import os

def generate_cam_document(company_data, risk_analysis):
    doc = Document()
    doc.add_heading('Credit Appraisal Memo', 0)
    
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(f"Company: {company_data.get('name', 'Unknown')}")
    
    doc.add_heading('Risk Analysis', level=1)
    doc.add_paragraph(f"Score: {risk_analysis.get('final_score')}")
    doc.add_paragraph(f"Decision: {risk_analysis.get('decision')}")
    
    os.makedirs('output', exist_ok=True)
    filepath = f"output/cam_{company_data.get('companyId', 'report')}.docx"
    doc.save(filepath)
    return filepath
