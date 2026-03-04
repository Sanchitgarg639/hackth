"""
Credit Appraisal Memo (CAM) Builder
Uses python-docx to generate an automated underwriter report.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import uuid
import logging
from datetime import datetime

logger = logging.getLogger(__name__)
OUTPUT_DIR = "/app/output"

def add_header(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def build_cam_document(data: dict) -> str:
    """
    Generates a structured DOCX Credit Appraisal Memo based on Phase 4 data.
    Returns the file path.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('CREDIT APPRAISAL MEMO (CAM)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated via Intelli-Credit Risk Engine — {datetime.now().strftime('%d %B %Y')}")
    
    company = data.get('companyData', {})
    risk = data.get('riskAnalysis', {})
    extracted = data.get('extractedData', {})
    ratios = extracted.get('ratios', {})
    findings = data.get('researchFindings', {})
    qualitative = data.get('qualitativeAssessment', {})
    
    # ── 1. Executive Summary ───────────────────────────
    add_header(doc, "1. Executive Summary", level=1)
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = 'Table Grid'
    
    def add_row(tab, r, label, val):
        tab.cell(r, 0).text = label
        tab.cell(r, 0).paragraphs[0].runs[0].font.bold = True
        tab.cell(r, 1).text = str(val)

    add_row(summary_table, 0, "Company Name", company.get('name', 'N/A'))
    add_row(summary_table, 1, "Sector", company.get('sector', 'N/A'))
    add_row(summary_table, 2, "Final Score", f"{risk.get('score', 'N/A')} / 100")
    add_row(summary_table, 3, "Probability of Default", f"{risk.get('pd', 0)*100:.2f}%")
    add_row(summary_table, 4, "Assigned Grade", risk.get('grade', 'N/A'))
    add_row(summary_table, 5, "Recommendation", risk.get('recommendation', 'N/A'))

    doc.add_paragraph()

    # ── 2. Financial Metrics (Capacity) ────────────────
    add_header(doc, "2. Financial Overview", level=1)
    fin_table = doc.add_table(rows=4, cols=2)
    fin_table.style = 'Light Shading Accent 1'
    
    add_row(fin_table, 0, "Requested Limit", f"INR {risk.get('recommendedLimit', 0):,.2f}")
    add_row(fin_table, 1, "Debt-to-Equity", f"{ratios.get('debtEquity', 'N/A')}x")
    add_row(fin_table, 2, "DSCR", f"{ratios.get('dscr', 'N/A')}x")
    add_row(fin_table, 3, "Cross-Verification Variance", f"{extracted.get('crossVerification', {}).get('variancePercent', 'N/A')}")
    
    doc.add_paragraph()
    
    # ── 3. Risk Engine Explainability (Character & Context) 
    add_header(doc, "3. ML Risk Driver Analysis", level=1)
    doc.add_paragraph("The XGBoost Probability of Default model identified the following key drivers using SHAP values:")
    
    reasons = risk.get('reasons', [])
    for rec in reasons:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"[{rec.get('impact', 'Neutral')}] ").bold = True
        p.add_run(rec.get('text', ''))

    doc.add_paragraph()
    
    # ── 4. Qualitative Assessment ──────────────────────
    add_header(doc, "4. Qualitative Remarks", level=1)
    doc.add_paragraph(f"Site Visit Rating: {qualitative.get('siteVisitRating', 'Not Rated')}/5")
    doc.add_paragraph(f"Management Quality: {qualitative.get('managementQualityRating', 'Not Rated')}/5")
    doc.add_paragraph("Credit Officer Notes:")
    doc.add_paragraph(qualitative.get('notes', 'No notes provided.'), style='Quote')
    
    doc.add_paragraph()
    
    # ── 5. Terms & Conditions ──────────────────────────
    add_header(doc, "5. Proposed Terms", level=1)
    doc.add_paragraph(f"Facility: Working Capital")
    doc.add_paragraph(f"Limit: INR {risk.get('recommendedLimit', 0):,.2f}")
    doc.add_paragraph(f"Interest Rate: {risk.get('suggestedInterestRate', 'N/A')}")
    doc.add_paragraph(f"Expected Loss (Provision): INR {risk.get('expected_loss', 0):,.2f}")
    
    doc.add_paragraph("\nDisclaimer: This is an automatically generated AI summary and requires formal sanctioning authority signatures.")

    # Save DOCX
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    file_id = f"CAM_{company.get('name', 'N_A').replace(' ', '_')}_{str(uuid.uuid4())[:8]}.docx"
    file_path = os.path.join(OUTPUT_DIR, file_id)
    doc.save(file_path)
    
    return file_id
